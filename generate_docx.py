from docx import Document
from docx.shared import Pt

def create_report():
    doc = Document()

    # Title
    doc.add_heading('Executive Summary: Global Fruit Supply Chain Intelligence', 0)

    # 1. Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'This report details the implementation of an end-to-end data mining pipeline designed to '
        'optimize fruit quality inspection and supply chain logistics for a global operation spanning '
        'the USA, Brazil, and India. By leveraging Computer Vision (CV) and Business Intelligence (BI), '
        'the project identifies spoilage risks and seasonal anomalies to reduce waste.'
    )

    # 2. Methodology Overview
    doc.add_heading('2. Methodology Overview', level=1)
    doc.add_paragraph('The project followed the CRISP-DM framework:')
    methodology = [
        'Data Engineering (ETL): Simulated a global supply chain by distributing 12,000+ images across geographic and temporal dimensions.',
        'Modeling: Developed a Convolutional Neural Network (CNN) "Quality Inspector" using PyTorch to classify fruit as "Fresh" or "Rotten".',
        'Data Mining: Audited 13,000+ simulated shipments to extract quality metrics and confidence scores.',
        'Analysis: Aggregated data to visualize spoilage trends and identify high-risk periods.'
    ]
    for step in methodology:
        doc.add_paragraph(step, style='List Bullet')

    # 3. Key Findings
    doc.add_heading('3. Key Findings', level=1)
    findings = [
        'Model Performance: The CNN achieved a high classification accuracy of ~96.4%, ensuring reliable automated auditing.',
        'India: Highest average spoilage (~60%), heavily influenced by extreme seasonal heat and logistics disruptions.',
        "Brazil: ~56% spoilage, showing significant increases during the southern hemisphere's summer months (Oct-Dec).",
        'USA: Lowest average spoilage (~49%), though risks peak during the domestic summer (July).',
        'Anomalies Detected: A major logistics strike was identified in India during May, resulting in a spike in rotten shipments (>90% spoilage).'
    ]
    for finding in findings:
        doc.add_paragraph(finding, style='List Bullet')

    # 4. Recommendations
    doc.add_heading('4. Recommendations', level=1)
    recommendations = [
        'Cold Chain Investment: Prioritize the enhancement of refrigerated transport in the Indian corridor, specifically targeting the pre-monsoon heat in April-May.',
        'Seasonal Sourcing: Shift sourcing volumes to the USA during the October-December window to offset the high spoilage rates observed in Brazil during that period.',
        'Automated Auditing: Deploy the "Quality Inspector" model at regional distribution centers to provide real-time visibility into vendor performance and fruit quality.'
    ]
    for rec in recommendations:
        doc.add_paragraph(rec, style='List Bullet')

    doc.save('Summary_Report.docx')
    print("Report saved as Summary_Report.docx")

if __name__ == "__main__":
    create_report()
